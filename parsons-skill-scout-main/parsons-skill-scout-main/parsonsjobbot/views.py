#Django imports
from django.views import generic
from django.urls import reverse
from django.shortcuts import redirect, render, get_object_or_404
from django.contrib import messages
from django.contrib.auth.mixins import LoginRequiredMixin
from django.views import View
from django.contrib.auth.models import Group, User
from django.db.models import Q
from functools import reduce

#imports 
from pyresparser import ResumeParser
from docx import Document
import spacy
from spacy.lang.en.stop_words import STOP_WORDS
import pandas as pd
import re
import tabula
import pandas as pd
import PyPDF2
import random
import matplotlib.pyplot as plt
import plotly.graph_objects as go
import plotly.express as px
import io
import urllib, base64

nlp = spacy.load('en_core_web_sm')

#imported models here

from .models import Candidate, UploadedFile, JobSubmission, Skill, UploadedXlsx, XlsxJob, SOWPosition, UploadedSOW, SimilarityScoreMatcher
from .forms import ResumeUploadForm, XlsxUploadForm, SOWUploadForm, UserSelectionForm


# Create your views here.


class HomeView(generic.TemplateView):
    """
    This view renders the home page for Skill Scout, but it does not have any major importance for the backend.

    Attributes:
        template_name (str): The name of the template used to render the home page.
        context_object_name (str): The name used to refer to the context data in the template.

    Methods:
        get_context_data(**kwargs): Override this method to add extra context data to the template.

    Usage:
        The HomeView is responsible for rendering the home page of Skill Scout. It extends the generic.TemplateView
        class provided by Django. The template_name attribute is set to "parsonsjobbot/home.html", which is the HTML
        template used to render the home page. The context_object_name attribute is set to "job", which means that the
        context data passed to the template will be accessible using the variable name "job".
    """
    template_name = "parsonsjobbot/home.html"
    context_object_name = "job"

class ResumeView(generic.FormView):
    """
    This view is responsible for handling the resume submission page. It allows users to submit their resumes, which are then
    parsed and associated with their candidate profiles through a one-to-one field with the user or can be a floating resume as well with no user assigned to it yet.

    Attributes:
        template_name (str): The name of the template used to render the resume submission page.
        form_class (Form): The form class used to handle the resume submission data.

    Methods:
        form_valid(form): Override this method to process the form data upon successful submission.

    Usage:
        The ResumeView is a subclass of generic.FormView provided by Django. It handles the resume submission process,
        allowing users to submit their resumes through a form. Upon submission, the form data is validated, and if valid,
        the form_valid() method is called to process the data. Otherwise, it gives an error to the user to fix.
    """
    template_name = 'parsonsjobbot/resumesubmissionpage.html'
    form_class = ResumeUploadForm


    def form_valid(self, form):
            """
            Handles the information from the HTML file and combines that to make a Candidate or update a candidate
            """
            try:
                resume_file = form.cleaned_data['resume']
                is_own_resume = form.cleaned_data['is_own_resume']
                resume_path = form.files['resume'].temporary_file_path() 

                # upload to admin backend
                uploaded_file = UploadedFile(file=resume_file)
            
                # parse the uploaded resume using pyresparser
                data = ResumeParser(resume_path).get_extracted_data()

                name = data['name']
                # for skills you can always take an input list and then add it to the skills found in the resume as well
                skills = ', '.join(data.get('skills', []))
                years_of_experience = data.get('total_experience', 0)
                try:
                    education = ', '.join(data.get('degree', []))
                except:
                    education = None

                # manual - if the parser doesnt catch everything
                if name != form.cleaned_data['name']:
                    name = form.cleaned_data['name']
                
                if not education:
                    if form.cleaned_data['education']:
                        education = form.cleaned_data['education']
                    
                if years_of_experience == 0:
                    if form.cleaned_data['years_of_experience']:
                        years_of_experience = form.cleaned_data['years_of_experience']
                        #print(years_of_experience)
                
                


                # if the user has already made a Candidate just update that candidate
                # visit the Candidate model for more information
                try:
                    candidate = Candidate.objects.get(user=self.request.user)
                    # if candidate wants to upload a resume on their behalf
                    if is_own_resume:
                        candidate.name=name
                        candidate.skills=skills
                        candidate.years_of_experience=years_of_experience
                        candidate.education=education

                        uploaded_file.save()

                        candidate.resume = uploaded_file
                        
                        candidate.save()

                        messages.success(self.request, "Resume Re-Upload Successful")
                        return redirect(reverse('parsonsjobbot:resume-home'))
                    
                    # if candidate wants to upload resume on another persons behalf
                    else:
                        candidate = Candidate.objects.create(
                            name=name,
                            skills=skills,
                            years_of_experience=years_of_experience,
                            education=education,
                        )

                        uploaded_file.save()

                        candidate.resume = uploaded_file
                        
                        candidate.save()

                        messages.success(self.request, "Resume Upload for another Candidate Successful")
                        return redirect(reverse('parsonsjobbot:resume-home'))

                # if the user has no candidate then make one for the user.
                except Candidate.DoesNotExist:

                    candidate = Candidate.objects.create(
                            user=self.request.user,
                            name=name,
                            skills=skills,
                            years_of_experience=years_of_experience,
                            education=education,
                        )
                    
                    uploaded_file.save()

                    candidate.resume = uploaded_file
                    
                    candidate.save()

                    messages.success(self.request, "Resume Uploaded Successfully")
                    return redirect(reverse('parsonsjobbot:resume-home'))
            except:
                messages.success(self.request, "Resume Upload unsuccessful")
                return redirect(reverse('parsonsjobbot:resume-home'))

class MetricsView(generic.TemplateView):
    """
    This view is responsible for displaying relevant skills and costs for the company, allowing stakeholders to track trends
    and focus areas.

    Attributes:
        template_name (str): The name of the template used to render the metrics page.
        context_object_name (str): The name used to refer to the context data in the template.

    Methods:
        get_context_data(**kwargs): Override this method to add extra context data to the template.

    Usage:
        The MetricsView is a subclass of generic.TemplateView provided by Django. It is responsible for rendering the
        metrics page, where stakeholders can view and analyze relevant data about the company's skills and costs. The view
        generates both a bar chart and a pie chart to represent the data visually.
    """
    template_name = "parsonsjobbot/metrics.html"
    context_object_name = "metrics"

    def get_context_data(self, **kwargs):
        context = super().get_context_data(**kwargs)

        skills = ['M1', 'M2', 'M3', 'M4', 'M5']
        num_jobs = [random.randint(50, 100) for _ in range(len(skills))]

        plt.figure(figsize=(10, 6))
        plt.bar(skills, num_jobs)
        plt.xlabel('Workflow Progression')
        plt.ylabel('Time')
        plt.title('Workflow progression vs Time')
        plt.xticks(rotation=45)
        
        buffer = io.BytesIO()
        plt.savefig(buffer, format='png')
        buffer.seek(0)
        image_png = buffer.getvalue()
        buffer.close()
        context['bar_chart'] = base64.b64encode(image_png).decode()

        categories = ['Sourcing', 'Screening', 'Interviewing', 'Offer', 'Onboarding']
        num_candidates = [random.randint(10, 30) for _ in range(len(categories))]

        fig = go.Figure(data=[go.Pie(labels=categories, values=num_candidates)])
        pie_chart = fig.to_html(full_html=False)

        context['pie_chart'] = pie_chart

        return context

class CandidateView(generic.ListView):
    """
    This view displays a list of candidates based on the user's group membership. Different groups have different access
    privileges to view candidate information and jobs associated with them.

    Attributes:
        template_name (str): The name of the template used to render the candidate view page.
        context_object_name (str): The name used to refer to the context data in the template.
        model (Model): The model associated with this view. In this case, it is the Candidate model.

    Methods:
        get_queryset(): Override this method to customize the queryset of candidates based on the user's group.

    Usage:
        The CandidateView is a subclass of generic.ListView provided by Django. It displays a list of candidates with
        different filtering criteria based on the user's group membership. The view uses the Candidate model to retrieve
        the candidate data.
    """
    template_name = "parsonsjobbot/job_recommendation.html"
    context_object_name = "job_rec"
    model = Candidate

    def get_queryset(self):
        user = self.request.user
        group_names = [group.name for group in user.groups.all()]
        search_query = self.request.GET.get('q')

        if search_query:
            queryset = Candidate.objects.filter(
                Q(name__icontains=search_query),
                Q(user=user) if 'Employee' in group_names or 'External' in group_names else Q()
            )
        elif 'Employee' in group_names or 'External' in group_names:
            queryset = Candidate.objects.filter(user=user)
        elif 'TOPM' in group_names or 'TA' in group_names or 'Admin' in group_names or 'PM' in group_names:
            queryset = Candidate.objects.all().order_by('name')
        else:
            queryset = Candidate.objects.none()

        return queryset

class CandidateDetailView(generic.DetailView):
    """
    This view displays detailed information about a candidate, including their personal information and available job
    listings. It replaces the previous Job Detail View and combines both the matched jobs view and the candidate detail
    view into one.

    Attributes:
        model (Model): The model associated with this view. In this case, it is the Candidate model.
        template_name (str): The name of the template used to render the candidate detail page.
        context_object_name (str): The name used to refer to the context data in the template.

    Methods:
        get_context_data(**kwargs): Override this method to add extra context data to the template.
        post(request, *args, **kwargs): Handle the HTTP POST request.

    Usage:
        The CandidateDetailView is a subclass of generic.DetailView provided by Django. It displays detailed information
        about a specific candidate, including their personal details and the job listings associated with them.
    """
    model = Candidate
    template_name = 'parsonsjobbot/candidate_detail.html'
    context_object_name = 'candidate'

    def get_context_data(self, **kwargs):
        context = super().get_context_data(**kwargs)
        candidate_id = self.kwargs['pk']
        candidate = Candidate.objects.get(pk=candidate_id)
        matched_jobs = get_open_positions_for_candidate(candidate)

        context['matched_jobs'] = matched_jobs
        context['reassign_user_form'] = UserSelectionForm()
        return context

    def post(self, request, *args, **kwargs):
        form = UserSelectionForm(request.POST)

        if form.is_valid():
            # this seems to not be working (1/09/2024)
            user = form.cleaned_data['user']
            candidate_id = self.kwargs['pk']
            candidate = Candidate.objects.get(pk=candidate_id)
            candidate.user = user
            candidate.save()

            return redirect(reverse('parsonsjobbot:resume-home'))

        context = self.get_context_data(**kwargs)
        context['form'] = form
        return self.render_to_response(context)

class CandidateDetailMatchedView(generic.DetailView):
    """
    This view displays detailed information about a candidate, showing only their personal information rendered onto a page.

    Attributes:
        model (Model): The model associated with this view. In this case, it is the Candidate model.
        template_name (str): The name of the template used to render the candidate detail page.
        context_object_name (str): The name used to refer to the context data in the template.

    Usage:
        The CandidateDetailMatchedView is a subclass of generic.DetailView provided by Django. It displays detailed
        information about a specific candidate, including their personal details, but it focuses solely on rendering the
        candidate information on a page.
    """
    model = Candidate
    template_name = 'parsonsjobbot/candidate_detail_match.html'
    context_object_name = 'candidate'

class MyCandidateProfileView(generic.DetailView):
    """
    This view displays the candidate information for the logged-in user. It shows details of the candidate's profile based
    on the user's membership in the 'Employee' or 'External' group.

    Attributes:
        model (Model): The model associated with this view. In this case, it is the Candidate model.
        template_name (str): The name of the template used to render the candidate profile page.
        context_object_name (str): The name used to refer to the context data in the template.

    Methods:
        get_object(queryset=None): Override this method to fetch the candidate model based on the logged-in user.
        get_context_data(**kwargs): Override this method to add extra context data to the template.

    Usage:
        The MyCandidateProfileView is a subclass of generic.DetailView provided by Django. It displays the candidate
        information for the logged-in user, based on their membership in the 'Employee' or 'External' group.
    """
    model = Candidate
    template_name = 'parsonsjobbot/my_candidate_profile.html'
    context_object_name = 'candidate'

    def get_object(self, queryset=None):
        user = self.request.user
        try:
            return Candidate.objects.get(user=user)
        except Candidate.DoesNotExist:
            return None

    def get_context_data(self, **kwargs):
        context = super().get_context_data(**kwargs)
        candidate = self.object
        
        if candidate:
            matched_jobs = get_open_positions_for_candidate(candidate)
            context['matched_jobs'] = matched_jobs
        
        return context

class AboutView(generic.TemplateView):
    """
    This view displays generic information about Skill Scout to the user. It has no real backend impact and is used solely
    for rendering an "About" page.

    Attributes:
        template_name (str): The name of the template used to render the "About" page.
        context_object_name (str): The name used to refer to the context data in the template.

    Usage:
        The AboutView is a subclass of generic.TemplateView provided by Django. It displays generic information about
        Skill Scout, intended to be shown as an "About" page.
    """
    template_name = "parsonsjobbot/about.html"
    context_object_name = "about"

class ProfileView(generic.TemplateView):
    """
    This view renders the user's profile information, including their groups and name. It will also display the user's
    candidate class once implemented. Users can edit their candidate information in this view or through a separate
    ProfileEdit view.

    Attributes:
        template_name (str): The name of the template used to render the profile page.
        context_object_name (str): The name used to refer to the context data in the template.

    Methods:
        get_context_data(**kwargs): Override this method to add extra context data to the template.

    Usage:
        The ProfileView is a subclass of generic.TemplateView provided by Django. It displays information about the user's
        profile, such as their groups and name. Additionally, it will eventually show the user's candidate class, allowing
        them to edit that information either within this view or via a separate ProfileEdit view.
    """
    template_name = "parsonsjobbot/profile.html"
    context_object_name = "profile"


    def get_context_data(self, **kwargs):
        context = super().get_context_data(**kwargs)
        if self.request.user.is_authenticated:
            user = self.request.user
            context['user'] = user
            context['name'] = user.get_full_name()
            context['groups'] = user.groups.all()
        return context

class EditProfileView(LoginRequiredMixin, View):
    """
    This view allows authenticated users to update their profile information. Only admins or users belonging to a higher
    level group should be able to change the groups of a user.

    Attributes:
        template_name (str): The name of the template used to render the edit profile page.

    Methods:
        get(request): Handle the HTTP GET request to render the edit profile page.
        post(request): Handle the HTTP POST request to update the user's profile information.

    Usage:
        The EditProfileView is a view for authenticated users to update their profile information. It allows users to edit
        their name and select the groups they belong to. Only admins or users with higher-level group permissions can
        change the groups for a user.
    """
    template_name = "parsonsjobbot/edit_profile.html"


    def get(self, request):
        if self.request.user.is_authenticated:
            context = {
                'user': request.user,
                'groups': Group.objects.all()
            }
            return render(request, self.template_name, context)
        return render(request, self.template_name)


    def post(self, request):
        user = request.user
        name = request.POST.get('name')
        group_ids = request.POST.getlist('groups')


        # Update name
        user.first_name, user.last_name = name.split(" ", 1) if " " in name else (name, "")
        user.save()


        # Update groups
        user.groups.set(group_ids)


        return redirect('parsonsjobbot:profile')

class JobSubmissionView(generic.CreateView):
    """
    This view allows users to submit job information using a form. The submitted data is saved to the backend.
    Eventually, it should support uploading an excel sheet to parse and extract data.

    Attributes:
        model (Model): The model associated with this view. In this case, it is the JobSubmission model.
        fields (list): The fields to include in the form for job submission.
        template_name (str): The name of the template used to render the job submission page.

    Methods:
        get_success_url(): Override this method to specify the URL to redirect to after a successful form submission.
        form_valid(form): Override this method to save the job submission data and associate it with skills.

    Usage:
        The JobSubmissionView is a subclass of generic.CreateView provided by Django. It allows users to submit job
        information using a form. The form fields include 'position_id', 'position_description', 'skill_level', and
        'job_title'. After successful form submission, the data is saved in the backend and associated with relevant skills.
    """

    model = JobSubmission
    fields = ['position_id', 'position_description', 'skill_level', 'job_title']
    template_name = 'parsonsjobbot/job_submission.html'

    def get_success_url(self):
        return reverse('parsonsjobbot:job-submission')

    def form_valid(self, form):
        job_submission = form.save(commit=False)
        job_submission.save()
        skills_input = self.request.POST.get('skills', '')
        skills = [skill.strip() for skill in skills_input.split(',') if skill.strip()]
        for skill_name in skills:
            skill, _ = Skill.objects.get_or_create(name=skill_name)
            job_submission.skills.add(skill)
        messages.success(self.request, "Job Uploaded Successfully")
        return redirect(self.get_success_url())

class XlsxSubmissionView(generic.FormView):
    """
    This view is designed to parse job information from NEE files sent by managers.
    It reads the NEE files, goes through each line, and extracts important information to store in the backend.
    The parsed data can be used for querying SOWs (Statements of Work) now.

    Attributes:
        template_name (str): The name of the template used to render the xlsx submission page.
        form_class (Form): The form class associated with this view. In this case, it is XlsxUploadForm.

    Methods:
        form_valid(form): Override this method to handle the form submission when it is valid.

    Usage:
        The XlsxSubmissionView is a subclass of generic.FormView provided by Django. It allows users to submit NEE files
        containing job information. The view then reads the NEE files, extracts the relevant data, and stores it in the
        backend. The form used in this view is XlsxUploadForm.
    """
    template_name = 'parsonsjobbot/xlsx_submission_page.html'

    form_class = XlsxUploadForm

    def form_valid(self, form):
        try:
            xlsx_file = form.cleaned_data['xlsx']

            xlsx_path = form.files['xlsx'].temporary_file_path()

            uploaded_xlsx_file = UploadedXlsx(file=xlsx_file)

            positions = extract_positions_from_excel(xlsx_path)
            try:
                upload_counter = 0
                for position in positions:
                    #check if there is a new file
                    xlsx_job = XlsxJob(**position)
                    xlsx_job_stored = XlsxJob.objects.filter(  
                        tonum = position['tonum'],
                        posnum = position['posnum'],
                        pdnum = position['pdnum'],
                        previous_names = position['previous_names'],
                        project = position['project'],
                        status = position['status'],
                        labor_cat = position['labor_cat'],
                        level = position['level'],
                        clin = position['clin'],
                        location = position['location'],
                        release_date = position['release_date'],
                        open_or_closed = position['open_or_closed']
                    )
                    if xlsx_job_stored.exists():
                        upload_counter += 1
                        continue
                    else:
                        xlsx_job = XlsxJob(**position)
                        xlsx_job.save()
                    
                if upload_counter != len(positions):
                    uploaded_xlsx_file.save()

            except:
                messages.success(self.request, "Something went wrong while parsing valid positions.")
                return redirect(reverse('parsonsjobbot:job-xlsx-submission'))
            
            messages.success(self.request, "Xlsx Upload Successful")
            return redirect(reverse('parsonsjobbot:job-xlsx-submission'))
        
        except:
            messages.success(self.request, "Xlsx Upload Unsuccessful, please orient it so that it matches a previously uploaded xlsx file")
            return redirect(reverse('parsonsjobbot:job-xlsx-submission'))

class SOWSubmissionView(generic.FormView):
    """
    This view allows users to submit Statement of Work (SOW) files containing position information.
    It reads the SOW files, parses tables and position descriptions, and stores the extracted data in the backend.

    Attributes:
        template_name (str): The name of the template used to render the SOW submission page.
        form_class (Form): The form class associated with this view. In this case, it is SOWUploadForm.

    Methods:
        form_valid(form): Override this method to handle the form submission when it is valid.

    Usage:
        The SOWSubmissionView is a subclass of generic.FormView provided by Django. It allows users to submit SOW files
        containing position information. The view then reads the SOW files, parses tables and position descriptions, and
        stores the extracted data in the backend. The form used in this view is SOWUploadForm.
    """
    template_name = 'parsonsjobbot/sow_submission_page.html'

    form_class = SOWUploadForm

    def form_valid(self, form):
        try:
            sow_file = form.cleaned_data['sow']

            sow_path = form.files['sow'].temporary_file_path()

            uploaded_sow_file = UploadedSOW(file=sow_file)

            if "pdf" in sow_file.name:
                positions = parse_tables_and_position_descs_pdf(sow_path)
            elif "docx" in sow_file.name:
                positions = parse_tables_and_position_descs_word(sow_path)
            else:
                messages.success(self.request, "SOW Upload Type is not PDF or Word (.docx)")
                return redirect(reverse('parsonsjobbot:sow-submission'))

            
            # important parts of positions looks like this for the most part:
            # positions = [
            #     {'Position ID': str,
            #      'Position Number': str,
            #      'Position Description': [str(position number), str(position title), str(position description)],
            #      'Skill Level': str,
            #      'Service Category': str,
            #      'Job Title': str,
            #      'Task Order Number': str},
            #      {
            #         ...
            #       },
            #       etc
            # ]

            total_pos = counter = len(positions)
            for position in positions:
                

                try:
                    sow_pos = SOWPosition.objects.get(pos_id=position['Position ID'], posnum=position['Position Number'])

                    sow_pos.tonum = position['Task Order Number']
                    sow_pos.pos_id = position['Position ID']
                    sow_pos.posnum = position['Position Number']
                    sow_pos.location = position['Location']
                    sow_pos.posdescnum = position['Position Description'][0]
                    sow_pos.posdesctitle = position['Position Description'][1]
                    sow_pos.posdesc = position['Position Description'][2]
                    sow_pos.level = position['Skill Level']
                    sow_pos.service_cat = position['Service Category']
                    sow_pos.job_title = position['Job Title']
                    sow_pos.save()

                    

                except SOWPosition.DoesNotExist:
                    sow_pos = SOWPosition.objects.create(
                        tonum = position['Task Order Number'],
                        pos_id = position['Position ID'],
                        posnum = position['Position Number'],
                        location = position['Location'],
                        posdescnum = position['Position Description'][0],
                        posdesctitle = position['Position Description'][1],
                        posdesc = position['Position Description'][2],
                        level = position['Skill Level'],
                        service_cat = position['Service Category'],
                        job_title = position['Job Title'],
                    )
                    sow_pos.save()

                    counter -= 1
            
            if counter != total_pos:
                uploaded_sow_file.save()

            messages.success(self.request, "SOW Upload Successful")
            return redirect(reverse('parsonsjobbot:sow-submission'))

        except KeyError:
            messages.success(self.request, "SOW Upload Unsuccessful - Something went wrong with making the position")
            return redirect(reverse('parsonsjobbot:sow-submission'))
        except:
            messages.success(self.request, "SOW Upload Unsuccessful - Something went wrong with parsing the SOW")
            return redirect(reverse('parsonsjobbot:sow-submission'))

class XlsxSOWOpenMatcherView(generic.ListView):
    """
    This view retrieves SOWPosition models related to open XlsxJob models based on matching 'tonum' and 'posnum'.
    It displays a list of SOWPosition objects associated with open XlsxJobs in the template.

    Attributes:
        template_name (str): The name of the template used to render the xlsx SOW open matcher page.
        context_object_name (str): The name of the context variable to use in the template for the list of SOWPosition objects.

    Methods:
        get_queryset(): Override this method to retrieve the queryset of SOWPosition objects related to open XlsxJob models.

    Usage:
        The XlsxSOWOpenMatcherView is a subclass of generic.ListView provided by Django. It retrieves SOWPosition models
        related to open XlsxJob models based on matching 'tonum' and 'posnum'. The retrieved SOWPosition objects are then
        displayed in the template using the context variable 'sow_positions'.
    """
    template_name = 'parsonsjobbot/xlsx_sow_open_matcher.html'
    context_object_name = 'sow_positions'

    def get_queryset(self):
        tonum_query = self.request.GET.get('tonum')
        posnum_query = self.request.GET.get('posnum')

        sow_positions = []

        if tonum_query and posnum_query:
            try:
                xlsxposition = XlsxJob.objects.filter(tonum=tonum_query, posnum=posnum_query, open_or_closed='open')
                sow_positions = SOWPosition.objects.filter(tonum=tonum_query, posnum=posnum_query)
                return sow_positions
            except XlsxJob.DoesNotExist:
                return SOWPosition.objects.none()
        elif tonum_query:
            try:
                xlsxposition = XlsxJob.objects.filter(tonum=tonum_query, open_or_closed='open')
                sow_positions = SOWPosition.objects.filter(tonum=tonum_query)
                return sow_positions
            except XlsxJob.DoesNotExist:
                return SOWPosition.objects.none()
        elif posnum_query:
            try:
                xlsxposition = XlsxJob.objects.filter(posnum=posnum_query, open_or_closed='open')
                sow_positions = SOWPosition.objects.filter(posnum=posnum_query)
                return sow_positions
            except XlsxJob.DoesNotExist:
                return SOWPosition.objects.none()
        # if no search parameters provided, show all open positions
        else:
            open_xlsx_jobs = XlsxJob.objects.filter(open_or_closed='open')
            tonum_list = open_xlsx_jobs.values_list('tonum', flat=True)
            posnum_list = open_xlsx_jobs.values_list('posnum', flat=True)

            tonum_posnum_pairs = zip(tonum_list, posnum_list)
            
            sow_positions = SOWPosition.objects.filter(reduce(lambda x, y: x | y, [Q(tonum=tonum, posnum=posnum) for tonum, posnum in tonum_posnum_pairs]))

        return sow_positions

class XlsxSOWOpenDeatilView(generic.DetailView):
    """
    This view displays detailed information about an open SOWPosition model associated with an XlsxJob.

    Attributes:
        model (Model): The model associated with this view. In this case, it is the SOWPosition model.
        template_name (str): The name of the template used to render the detailed view of the open SOWPosition.
        context_object_name (str): The name of the context variable to use in the template for the SOWPosition object.

    Usage:
        The XlsxSOWOpenDeatilView is a subclass of generic.DetailView provided by Django. It displays detailed information
        about an open SOWPosition model associated with an XlsxJob. The view uses the SOWPosition model and the specified
        template to render the detailed information.
    """
    model = SOWPosition
    template_name = 'parsonsjobbot/xlsx_sow_open_detail.html'
    context_object_name = 'position_detail'

class XlsxSOWOpenDeatilMatchView(generic.DetailView):
    """
    This view displays detailed information about an open SOWPosition model associated with an XlsxJob, along with
    the candidates matched to the position.

    Attributes:
        model (Model): The model associated with this view. In this case, it is the SOWPosition model.
        template_name (str): The name of the template used to render the detailed view of the open SOWPosition with matched candidates.
        context_object_name (str): The name of the context variable to use in the template for the SOWPosition object.

    Methods:
        get_context_data(**kwargs): Override this method to add additional context data to be used in the template.

    Usage:
        The XlsxSOWOpenDeatilMatchView is a subclass of generic.DetailView provided by Django. It displays detailed information
        about an open SOWPosition model associated with an XlsxJob, along with the candidates matched to the position.
        The view uses the SOWPosition model and the specified template to render the detailed information along with the
        matched candidates.
    """
    model = SOWPosition
    template_name = 'parsonsjobbot/xlsx_sow_detail_match.html'
    context_object_name = 'position'

    def get_context_data(self, **kwargs):
        context = super().get_context_data(**kwargs)
        sow_position_id = self.kwargs['pk']
        sow_position = SOWPosition.objects.get(pk=sow_position_id)
        
        matched_candidates = get_candidate_for_position(sow_position)

        context['matched_candidates'] = matched_candidates
        return context

#helper functions / redirects to go back to other apps

def redirect_to_landing_page(request):
    """
    return to the other app landing_page
    """
    return redirect('landing_page:landing_page')

def reorganize_dict(positions_list: list[dict]) -> list[dict]:
    """
    Seprating the ToNum key in dictionary with the PosNum as well and then putting it in the dictonary independently separated
    Also chnaged the Date to look more realistic and not just a datetime object
    """
    for position in positions_list:
        tonum_parts = position['tonum'].split(' / ')
        position['tonum'] = tonum_parts[0].strip()
        posnum = tonum_parts[1].strip().replace('\n', ' ')
        if " " in posnum:
            position['posnum'] = posnum.split(' ')[0]
        else:
            position['posnum'] = posnum
        position['release_date'] = position['release_date'].strftime('%Y-%m-%d')
    return positions_list

def extract_positions_from_excel(xlsx_file: UploadedFile) -> list[dict]:
    """
    Interpret the NEE xlsx files and return a list of positions

    Args:
        xlsx_file (str): The path to the NEE xlsx file.

    Example return object (all key value pairs are strings):
    [
        [{'TONum': '1', 'PDNum': '9', 'Previous Names': 'TO FILL', 'Project': 'Requirements Engineering', 
        'Status': 'First come, first serve', 'Labor Cat': 'Systems Engineer', 'Level': '2 - Mid (6+ to 12)', 'CLIN': '1', 
        'Location': 'WMA-CS', 'Release Date': '2023-05-23', 'PosNum': '34'}, ...,

        {'TONum': '16', 'PDNum': 4, 'Previous Names': 'TO FILL', 'Project': 'CLOSED ', 
        'Status': 'First come, first serve', 'Labor Cat': 'Systems Engineer', 'Level': '3 - Senior (12+ to 18)', 'CLIN': '0009', 
        'Location': 'WMA-CS', 'Release Date': '2023-06-13', 'PosNum': '9'}, ...]
    ]
    """
    df = pd.read_excel(xlsx_file, header=None)
    open_positions = []
    closed_positions = []


    position_headers = ["tonum", "pdnum", "previous_names", "project", "status", "labor_cat", "level", "clin", "location", "release_date"]
    det = False
    for row_id, row in df.iterrows():
        if "open positions" in str(row[3]).lower():
            det = True
        elif "closed positions" in str(row[3]).lower():
            det = False

        # if we are in open position and we are not in the header row
        elif det and not ("project" in str(row[3]).lower()):
            zipped_dict_row = dict(zip(position_headers, row))
            zipped_dict_row["open_or_closed"] = 'open'
            open_positions.append(zipped_dict_row)
        elif not det and not pd.isna(row[3]):
            zipped_dict_row = dict(zip(position_headers, row))
            zipped_dict_row["open_or_closed"] = 'closed'
            closed_positions.append(zipped_dict_row)


    positions = reorganized_positions = reorganize_dict(open_positions)
    reorganized_closed_positions = reorganize_dict(closed_positions)

    positions.extend(reorganized_closed_positions)

    return positions

def parse_tables_and_position_descs_word(path: UploadedFile) -> list[dict]:
    """
    Works with SOW Task orders that are word documents... mostly.

    Args:
        path (str): The file path to the SOW Task order Word document.

    Returns:
        list: A list of dictionaries containing job position details extracted from the SOW.
    
    Example:
        positions = parse_tables_and_position_descs_word('path/to/sow_task_order.docx')
        # Output:
        # positions = [
        #     {'Task Order Number': 'TO-123', 'Position ID': 'PID-456', 'Position Number': '456',
        #      'Location': 'Location 1', 'Position Description': ['PD123', 'Position Title', 'Description of the position'],
        #      'Skill Level': '3 - Senior (12 to 18 years)', 'Service Category': 'Service Category 1', 'Job Title': 'Job Title 1'},
        #     ...
        # ]
    """

    def get_position_descriptions(doc):
        """
        Gather the Descriptions inside the SOW.

        Args:
            doc (python-docx.Document): The Word document object from which the position descriptions will be extracted.

        Returns:
            dict: A dictionary containing position descriptions from the SOW. The keys are the position names, and the values are the descriptions.
        """
        past_appendix_A_key_found = appendix_B_found = False
        current_position = 'placeholder'
        curr_position_dict = {current_position : ''}
        for paragraph in doc.paragraphs:
            if "appendixakey:" in ''.join(paragraph.text.split()).lower():
                past_appendix_A_key_found = True
            if "appendixb:positiondescriptions" in ''.join(paragraph.text.split()).lower() and past_appendix_A_key_found:
                appendix_B_found = True
            if appendix_B_found:
                # new position found using regex
                if re.search("^Position[s]{0,1} [0-9]+[a-z]*:", paragraph.text):
                    current_position = ' '.join(paragraph.text.split()).strip()
                    curr_position_dict[current_position] = ''
                else:
                    curr_position_dict[current_position] += ' '.join(paragraph.text.split()).strip() + ' '

        return curr_position_dict

    def find_table_count(text):
        """
        Find the right tables we want in the backend from the SOWs, all tables should have some variation of "Position".

        Args:
            text (python-docx.Document): The Word document object containing the SOW text.

        Returns:
            list: A list of table indices (integers) where each table contains a cell with the exact text "Position ID".
        """
        tables = []
        table_count = 0
        for table in text.tables:
            for row in table.rows:
                for cell in row.cells:
                    if cell.text.strip("*") == "Position ID":
                        tables.append(table_count)
            #increment the table count to match the proper table we want
            table_count += 1
        return tables
    
    

    doc = Document(path)
    # get the positions in the SOWs, preprocess the keys
    position_desc_dict = get_position_descriptions(doc)
    position_desc_keys = position_desc_dict.keys()

    table_seeked = find_table_count(doc)
    total_jobs = []

    for table in table_seeked:
        cur_table = doc.tables[table]

        # preprocess the labels
        cur_row_label = [cell.text.strip("*").strip(" \n#") for cell in cur_table.rows[0].cells]

        for row in cur_table.rows[1:]:  # Exclude the header row
            # zip together the dictionary with the labels and the row values
            cur_row_creation = dict(zip(cur_row_label, [cell.text.lstrip('0').strip('\n') for cell in row.cells]))

            # Get the task order number
            cur_pos_id = cur_row_creation.get('Position ID')
            tonum = cur_pos_id.split('-')[1].lstrip('0')
            cur_row_creation['Task Order Number'] = tonum

            # process the position description to get the 3 parts:
            # [position description number, position description title, position description]
            position_desc = cur_row_creation.get("Position Description")
            for key in position_desc_keys:
                if position_desc in key:
                    cur_row_creation['Position Description'] = [position_desc, key, position_desc_dict[key]]

            # check if job has a position number:
            pos_num = cur_row_creation.get('Position Number')
            if not pos_num:
                try:
                    cur_row_creation['Position Number'] = cur_pos_id.split('-')[-1].lstrip('0')
                except:
                    print('Position Number creation from the position id failed')
            #print(cur_row_creation)
            total_jobs.append(cur_row_creation)
    
    # fill in missing values with keys we will need for processing data
    keys = ['Task Order Number', 'Position ID', 'Position Number', 'Location', 'Position Description', 'Skill Level', 'Service Category', 'Job Title']
    experience_levels = ['1 - Junior (0 to 6 years)', '2 - Mid (6 to 12 years)', '3 - Senior (12 to 18 years)', '4 - Expert (18 or more years)']

    for position in total_jobs:
        # fill in missing values with N/A
        for key in keys:
            if key not in position.keys():
                position[key] = 'N/A'
        # fill in the missing values for the list in Position Description
        if type(position['Position Description']) is not list and type(position['Position Description']) is str:
            position['Position Description'] = [position['Position Description'], 'N/A', 'N/A']
        # organize the Skill Level so its more streamlined
        for exp_level in experience_levels:
            if exp_level.split(' - ')[0] in position['Skill Level']:
                position['Skill Level'] = exp_level

    return total_jobs

def parse_tables_and_position_descs_pdf(path: UploadedFile) -> list[dict]:
    """
    Works with SOW Task orders that are PDF documents... mostly (does not include OCR).

    Args:
        path (str): The file path to the SOW Task order PDF document.

    Returns:
        list: A list of dictionaries containing job position details extracted from the SOW.
    
    Example:
        positions = parse_tables_and_position_descs_pdf('path/to/sow_task_order.pdf')
        # Output:
        # positions = [
        #     {'Task Order Number': 'TO-123', 'Position ID': 'PID-456', 'Position Number': '456',
        #      'Location': 'Location 1', 'Position Description': ['PD123', 'Position Title', 'Description of the position'],
        #      'Skill Level': '3 - Senior (12 to 18 years)', 'Service Category': 'Service Category 1', 'Job Title': 'Job Title 1'},
        #     ...
        # ]
    """

    def process_document(path: str) -> list[dict]:
        """
        This function parses the pdf document to find the tables and job numbers
        Works mostly, except for sometimes with the "TO Section" and "Position Description" it sometimes is not found
        """

        encodings = ['utf-8', 'cp1252', 'latin-1']
        for encoding in encodings:
            try:
                tables = tabula.read_pdf(input_path=path, pages='all', multiple_tables=True, encoding=encoding)
                if tables is not None:
                    break
            except:
                continue
        
        total_data_rows = []

        for table in tables:
            header_row = None
            data_rows = []

            # Find the header row containing "Position ID*"
            for row_index, row in enumerate(table.itertuples(index=False), start=1):
                if "Position ID*" in row or "Position ID" in row:
                    header_row = row
                    break

            if header_row is None:
                continue  # Skip table if no header row found

            # Create dictionaries using header row values as keys
            for row_index, row in enumerate(table.itertuples(index=False), start=1):
                if row == header_row or pd.isna(row[header_row.index("Position ID*")]):
                    continue  # Skip header row and rows with NaN Position ID*

                # Create dictionary using header values as keys and row values as values
                dictionary = {key: value for key, value in zip(header_row, row) if not pd.isna(value)}
                if "Position ID*" in dictionary:
                    dictionary["Position ID"] = dictionary.pop("Position ID*")
                data_rows.append(dictionary)

            total_data_rows.extend(data_rows)
        
        return total_data_rows

    def get_position_descs(path: str) -> dict:
        """
        This function gets the position descriptions in the SOWs in appendix B
        """

        start_page = -1
        positions = {}

        with open(path, 'rb') as file:
            reader = PyPDF2.PdfReader(file)

            #Iterate through each page
            for page_number in range(len(reader.pages)):
                page = reader.pages[page_number]
                text = page.extract_text()

                #Check if the page contains "Appendix B"
                if 'Appendix B' in text and "Overall Assignment Description" in text:
                    #print(text)
                    start_page = page_number + 1
                    break

            if start_page != -1:
                #print(text)
                total_pages = len(reader.pages)
                # print(f"Beginning page number: {start_page}")
                # print(f"Total number of pages: {total_pages}")
            else:
                print("Error: Appendix B not found in the PDF document.")
            pos_desc_num = 0
            for page_number in range(start_page-1, total_pages):
                
                page = reader.pages[page_number]
                text = page.extract_text()
                #filter the list before processing
                filtered_list = filter(text)
                if any(item.startswith("Position") for item in filtered_list):
                    pos_desc_num += 1
                    positions[pos_desc_num] = []
                positions[pos_desc_num].extend(filtered_list)
            #combine the list into sentences
            for position_key in positions:
                new_desc = []
                curr_sentence = ''
                for element in positions[position_key]:
                    curr_sentence += element + ' '
                    if '.' in element:
                        new_desc.append(curr_sentence.strip())
                        curr_sentence = ''
                positions[position_key] = new_desc
            #separate each sentence field now in terms of the headers such as Position #: Title, Overall Assignment Description: Description, etc
            # for desc in positions.values():
            #     print(desc)
            #     print()

            return positions

    def filter(text) -> list:
        formating_text = text.split("\n")
        filtered_list = [(' '.join(item.replace('UNCLASSIFIED', '').strip().replace('\uf0b7', '').split())).lstrip('o ') 
                        for item in formating_text if item.strip() != '' and "Page" not in item]
        if 'Appendix B: Position Descriptions' in filtered_list:
            filtered_list.remove('Appendix B: Position Descriptions')
        while '' in filtered_list:
            filtered_list.remove('')
        return filtered_list

    positionIDs = process_document(path)
    positionDescs = get_position_descs(path)
    
    newPositions = []
    keys = ['Task Order Number', 'Position ID', 
            'Position Number', 'Location', 
            'Position Description', 'Skill Level', 
            'Service Category', 'Job Title']
    experience_levels = ['1 - Junior (0 to 6 years)', '2 - Mid (6 to 12 years)', 
                         '3 - Senior (12 to 18 years)', '4 - Expert (18 or more years)']

    for position in positionIDs:
        # get the position number and description, then the title
        currPos = {}
        cur_pos_id = position.get('Position ID', 'N/A')
        if cur_pos_id:
            tonum = cur_pos_id.split('-')[1].lstrip('0')
            currPos['Task Order Number'] = tonum
            currPos['Position ID'] = cur_pos_id
            currPos['Position Number'] = cur_pos_id.split('-')[-1].lstrip('0')
        else:
            currPos['Task Order Number'] = "N/A"
            currPos['Position ID'] = "N/A"
            currPos['Position Number'] = "N/A"
        currPos['Location'] = position.get('Location', 'N/A')
        # format the descriptor list
        posDescNum = position.get('Description', 'N/A')
        descriptorList = [posDescNum]
        # get the job title
        for positionNumber in positionDescs:
            if int(positionNumber) == int(posDescNum):
                text = " ".join(positionDescs[positionNumber])
                title, description = text.split(" Overall Assignment Description: ")[0], text.split("Overall Assignment Description:")[-1].strip()
                title = title.split(":")[-1].strip()
                descriptorList.append(title)
                descriptorList.append(description)
        currPos['Position Description'] = descriptorList
        # normalize Skill levels for filtering
        for skillLevel in experience_levels:
            level = skillLevel[0]
            if level == position.get('Skill Level')[0]:
                currPos['Skill Level'] = skillLevel
        currPos['Service Category'] = position.get('Service Category', 'N/A')
        currPos['Job Title'] = position.get('Job Title', 'N/A')
        
        #double check we have used all of our keys, put in a default value:
        if len(keys) == len(currPos):
            newPositions.append(currPos)
        else:
            for key in keys:
                if key not in currPos:
                    currPos[key] = "N/A"
            newPositions.append(currPos)
        
    return newPositions

def compute_similarity_percentage(str1: str, str2: str) -> float:
    """
    Computes the similarity percentage between two input strings using spaCy.

    Args:
        str1 (str): The first input string.
        str2 (str): The second input string.

    Returns:
        float: The similarity percentage between the two strings.

    Example:
        similarity = compute_similarity_percentage("This is a sample text.", "Sample text for comparison.")
        # Output: similarity = 88.20533968855472
    """

    def preprocess_text(text):
        nlp = spacy.load("en_core_web_sm")
        doc = nlp(text)

        tokens = [token.lemma_.lower() for token in doc if not token.is_stop and not token.is_punct and token.text.lower() not in STOP_WORDS]
        preprocessed_text = ' '.join(tokens)

        return preprocessed_text

    nlp = spacy.load("en_core_web_lg")
    with nlp.disable_pipes():
        doc1 = nlp(preprocess_text(str1))
        doc2 = nlp(preprocess_text(str2))
    
    similarity_percentage = doc1.similarity(doc2) * 100
    return similarity_percentage

def get_open_positions_for_candidate(candidate: Candidate) -> list[(SOWPosition, float)]:
    """
    Modularized getting open positions for a candidate so that it can move wherever needed easier.
    Also caches the processes so that it is stored for later use.

    Args:
        candidate (Candidate): The candidate model instance for which open positions are to be found.

    Returns:
        list: A list of tuples containing matched SOWPosition models and their similarity scores.

    Example:
        candidate = Candidate.objects.get(id=1)
        matched_jobs = get_open_positions_for_candidate(candidate)
        # Output: matched_jobs = [(<SOWPosition: SOWPosition object (1)>, 83.32475812503635), ...]
    """

    #print("in here")
    candidate_skills = candidate.skills
    candiate_skill_and_education = candidate_skills + ', ' + candidate.education
    # get all XlsxJob models
    sow_positions = []
    open_xlsx_jobs = XlsxJob.objects.filter(open_or_closed='open')
    candidate_similarity_scores = SimilarityScoreMatcher.objects.filter(candidate_name=candidate.name, candidate_info=candiate_skill_and_education)
    tonum_list = open_xlsx_jobs.values_list('tonum', flat=True)
    posnum_list = open_xlsx_jobs.values_list('posnum', flat=True)
    tonum_posnum_pairs = zip(tonum_list, posnum_list)
    
    sow_positions = SOWPosition.objects.filter(reduce(lambda x, y: x | y, [Q(tonum=tonum, posnum=posnum) for tonum, posnum in tonum_posnum_pairs]))
    #print(sow_positions)
    # gather the similarity score matching
    similarity_scores = []
    for sow_position in sow_positions:
        # its good that we cache but if we update the resume of a candidate then we have a problem that it will 
        # of course give the right score based on the candidates resume used but it wont clear space so we need to clear next
        cached = SimilarityScoreMatcher.objects.filter(candidate_name=candidate.name, candidate_info=candiate_skill_and_education, sow_info=sow_position.posdesc)
        if cached:
            similarity_scores.append(cached[0].similarity_score)
        else:
            #score calculation is very expensive so if we did it once already just store that for later use
            similarity_score_instance = compute_similarity_percentage(candiate_skill_and_education, sow_position.posdesc)
            SimilarityScoreMatcher.objects.create(
                candidate_name=candidate.name, 
                candidate_info=candiate_skill_and_education, 
                sow_info=sow_position.posdesc,
                similarity_score=similarity_score_instance
            )
            similarity_scores.append(similarity_score_instance)
    # setup a list of tuples of sow_positions and similarity scores
    matched_jobs = list(zip(sow_positions, similarity_scores))
    matched_jobs.sort(key=lambda job: job[1], reverse=True)
    return matched_jobs

def get_candidate_for_position(sow_position: SOWPosition) -> list[(Candidate, float)]:
    """
    Modularized getting candidates for a given SOWPosition based on similarity of skills and education.
    Also caches the processes so that it is stored for later use.

    Args:
        sow_position (SOWPosition): The SOWPosition object for which matching candidates are to be found.

    Returns:
        list: A list of tuples containing matched Candidate models and their similarity scores.

    Example:
        sow_position = SOWPosition.objects.get(id=1)
        matched_candidates = get_candidate_for_position(sow_position)
        # Output: matched_candidates = [(<Candidate: Candidate object (1)>, 83.32475812503635), ...]
    """


    #print("in here")
    sow_position_desc = sow_position.posdesc
    # get all Candidate models
    candidates = Candidate.objects.all()
    similarity_scores = []
    for candidate in candidates:
        candidate_skills = candidate.skills
        candiate_skill_and_education = candidate_skills + ', ' + candidate.education
        cached = SimilarityScoreMatcher.objects.filter(candidate_name=candidate.name, candidate_info=candiate_skill_and_education, sow_info=sow_position_desc)
        if cached:
            similarity_scores.append(cached[0].similarity_score)
        else:
            #score calculation is very expensive so if we did it once already just store that for later use
            similarity_score_instance = compute_similarity_percentage(candiate_skill_and_education, sow_position.posdesc)
            SimilarityScoreMatcher.objects.create(
                candidate_name=candidate.name, 
                candidate_info=candiate_skill_and_education, 
                sow_info=sow_position_desc,
                similarity_score=similarity_score_instance
            )
            similarity_scores.append(similarity_score_instance)
    # setup a list of tuples of sow_positions and similarity scores
    matched_candidates = list(zip(candidates, similarity_scores))
    matched_candidates.sort(key=lambda job: job[1], reverse=True)
    return matched_candidates


#Out of service Views and Helpers

def calculate_score(candidate: Candidate, job_submission: SOWPosition) -> float:
    """
    this is a function used to help assist in deermining the score for the MatchedJobsView.
    Assigns weight to a job in comparison to the candidate where the better score means a higher score
    """
    # Example scoring mechanism
    score = 0
   
    # Increase score based on matching skills
    candidate_skills = set([skill.strip().lower() for skill in candidate.skills.split(',')])
    job_skills = set([skill.name.lower() for skill in job_submission.skills.all()])
    matched_skills = candidate_skills.intersection(job_skills)
    num_matched_skills = len(matched_skills)
    total_skills_required = len(job_skills)


    # Calculate percentage match
    percentage_match = (num_matched_skills / total_skills_required) * 100
    score += percentage_match
   
    # Adjust score based on other criteria (e.g., skill level, position description)
    # Add relevant scoring logic based on your requirements


    return score

class JobListView(View):
    """
    Just a generic about skill scout view for the user, no real backend impact just a page render for now
    This should be a list view eventually
    """
    template_name = "parsonsjobbot/job_listing.html"
    context_object_name = "jobs"


    # def get_queryset(self): 
    #     # Replace <candidate_id> with the actual candidate ID 
    #     candidate = Candidate.objects.get(id=<candidate_id>) 
    #     # Assuming skills are comma-separated
    #     candidate_skills = candidate.skills.split(",")    
    #     # Filter jobs based on skills and years of experience
    #     jobs = Job.objects.filter(keywords__in=candidate_skills, years_of_experience_required__lte=candidate.years_of_experience)
    #     return jobs

    def get_context_data(self, **kwargs):
        """
        This will need a double query, one for the candidate, and one for the total jobs possible as well. 
        Just make another try except block for Jobs and do context['jobs'] instead of context['candidate']
        at least this is what I am thinking for now.
        """
        context = super().get_context_data(**kwargs)
        
        # Get the current user's candidate profile
        user = self.request.user
        try:
            candidate = Candidate.objects.get(user=user)
            context['candidate'] = candidate
        except Candidate.DoesNotExist:
            pass
        
        return context

class MatchedJobsView(generic.ListView):
    """
    This view queries both the candidate and the job descriptions.
    The filtered query should parse and match the candidates skills with the job skills but its not refined yet.
    """
    # model relates to Candidate, XlsxJob, and SOWPosition but no one of those models overwrite the others
    model = None

    template_name = 'parsonsjobbot/job_listing.html'
    context_object_name = 'matched_jobs'

    def get_queryset(self):

        def get_open_positions_for_candidate(candidate):
            """
            Modularized getting open positions for a candidate it so that it can move wherever needed easier
            """

            candidate_skills = candidate.skills

            # get all XlsxJob models
            xlsx_jobs = XlsxJob.objects.all()

            # create a list to store the related SOWPosition models
            sow_positions = []

            # iterate over each XlsxJob and retrieve related SOWPosition models
            for xlsx_job in xlsx_jobs:
                tonum = xlsx_job.tonum
                posnum = xlsx_job.posnum
                open_or_closed = xlsx_job.open_or_closed
                # query SOWPosition models with matching tonum and posnum
                if open_or_closed == 'open':
                    positions = SOWPosition.objects.filter(tonum=tonum, posnum=posnum)

                    # append the positions to the list
                    sow_positions.extend(positions)
            
            # gather the similarity score matching
            similarity_scores = []

            for sow_positition in sow_positions:
                similarity_scores.append(compute_similarity_percentage(candidate_skills, sow_positition.posdesc))
            
            # setup a list of tuples of sow_positions and similarity scores
            matched_jobs = list(zip(sow_positions, similarity_scores))

            matched_jobs.sort(key=lambda job: job[1], reverse=True)

            return matched_jobs
    
        #gather the candidate
        candidate_id = self.kwargs['pk']
        candidate = Candidate.objects.get(pk=candidate_id)
        
        matched_jobs = get_open_positions_for_candidate(candidate)
        

        # return a query of [(sow_position, similarity score)] in descending order
        return matched_jobs
        


    def get_context_data(self, **kwargs):
        context = super().get_context_data(**kwargs)
        candidate_id = self.kwargs['pk']
        context['candidate'] = Candidate.objects.get(pk=candidate_id)
        return context

class JobDetailView(generic.DetailView):
    """
    Formerly used class when I was just enumerating uploaded resume forms but this now is out of work
    """
    model = JobSubmission
    template_name = "parsonsjobbot/job_detail.html"
    context_object_name = "job_detail"